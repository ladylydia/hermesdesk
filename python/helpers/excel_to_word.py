"""Copy the first worksheet of an .xlsx/.xlsm into a simple Word document (workspace-local)."""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any


def run(args: dict[str, Any]) -> dict[str, Any]:
    rel = str(args.get("excel_path") or args.get("path") or "").strip()
    if not rel:
        return {"ok": False, "error": "excel_path is required (relative to workspace)"}

    ws = Path(os.environ.get("HERMESDESK_WORKSPACE", ".")).resolve()
    src = (ws / rel).resolve()
    try:
        src.relative_to(ws)
    except ValueError:
        return {"ok": False, "error": "excel_path must stay inside the workspace"}
    if not src.is_file():
        return {"ok": False, "error": f"not a file: {src}"}
    if src.suffix.lower() not in (".xlsx", ".xlsm"):
        return {"ok": False, "error": "only .xlsx / .xlsm are supported (not legacy .xls)"}

    out_arg = args.get("output_path")
    if out_arg:
        dst = (ws / str(out_arg).strip()).resolve()
        try:
            dst.relative_to(ws)
        except ValueError:
            return {"ok": False, "error": "output_path must stay inside the workspace"}
        if dst.suffix.lower() != ".docx":
            return {"ok": False, "error": "output_path should end with .docx"}
    else:
        dst = src.with_suffix(".docx")

    try:
        from openpyxl import load_workbook  # type: ignore
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - import guard
        return {
            "ok": False,
            "skipped": True,
            "reason": f"openpyxl / python-docx not available in this bundle: {exc}",
        }

    sheet_name = args.get("sheet")
    wb = load_workbook(src, read_only=True, data_only=True)
    try:
        names = wb.sheetnames
        if not names:
            return {"ok": False, "error": "workbook has no sheets"}
        if sheet_name:
            key = str(sheet_name)
            if key not in wb.sheetnames:
                return {"ok": False, "error": f"unknown sheet {key!r}; available: {names!r}"}
            sh = wb[key]
        else:
            sh = wb[names[0]]

        doc = Document()
        doc.add_heading(src.name, level=0)
        rows = 0
        for row in sh.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            line = " | ".join(c.strip() for c in cells).strip(" |")
            if line:
                doc.add_paragraph(line)
                rows += 1
    finally:
        wb.close()

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    return {
        "ok": True,
        "excel": str(src.relative_to(ws)),
        "docx": str(dst.relative_to(ws)),
        "rows_written": rows,
        "sheet": str(sh.title),
    }
